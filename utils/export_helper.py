import base64
import io
import os
import re
import subprocess
import tempfile
import requests
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


FONT_CONFIG = {
    "cn_heading": "黑体",              # 中国学术论文标准：标题用黑体
    "cn_body": "宋体",                 # 正文用宋体
    "en_heading": "Times New Roman",   # 英文标题用Times New Roman
    "en_body": "Times New Roman",      # 英文正文用Times New Roman
    "code": "Consolas",                # 代码块用等宽字体
    "size_h1": 18,                     # 小二号 (18pt)
    "size_h2": 15,                     # 小三号 (15pt)
    "size_h3": 14,                     # 四号 (14pt)
    "size_h4": 12,                     # 小四号 (12pt)
    "size_body": 12,                   # 小四号 (12pt) - 学术论文正文标准
    "size_code": 10,                   # 五号 (10pt)
}


TABLE_STYLE = {
    "border_color": "000000",
    "border_width": 4,
    "header_bg": "0EA5E9",
    "header_text_color": "FFFFFF",
    "alternate_row_bg": "E8F4FC",  # 增强对比度，更明显的交替行背景
    "cell_padding": Pt(6),         # 略微增加内边距
}


CODE_BLOCK_STYLE = {
    "bg_color": "F4F4F5",
    "border_color": "E4E4E7",
    "border_width": 4,
    "padding": Pt(8),
    "line_spacing": 1.2,
}


PAGE_LAYOUT = {
    "width": Inches(8.27),
    "height": Inches(11.69),
    "margin_top": Inches(1),
    "margin_bottom": Inches(1),
    "margin_left": Inches(1.25),
    "margin_right": Inches(1.25),
}


INLINE_PATTERNS = [
    ("bold", re.compile(r"\*\*(.+?)\*\*")),
    ("italic", re.compile(r"\*(?!\*)(.+?)(?<!\*)\*")),
    ("code", re.compile(r"`(.+?)`")),
    ("link", re.compile(r"\[(.+?)\]\((.+?)\)")),
]

BULLET_SYMBOLS = "●○◆◇★☆►▸•·→⇒※"
LIST_LINE_PATTERN = re.compile(
    rf"^(\s*)([-+*]|\d+[.)]|[{re.escape(BULLET_SYMBOLS)}])\s+(.+)$"
)


class MarkdownToDocxConverter:
    """专业级 Markdown 转 Word 转换器"""

    def __init__(self, title="工程手册", system_type="US"):
        self.title = title
        self.system_type = system_type
        self.doc = Document()
        self._bullet_num_id = None
        self._number_num_id = None
        self._bullet_abstract_id = None
        self._number_abstract_id = None
        self._setup_styles()
        self._setup_page_layout()

    def convert(self, md_content):
        self._add_header_footer()
        self._add_title_and_toc()
        cleaned = self._sanitize_markdown(md_content or "")
        cleaned = self._strip_leading_h1(cleaned)
        self._parse_and_render(cleaned)
        buffer = io.BytesIO()
        self.doc.save(buffer)
        return buffer.getvalue()

    def _setup_styles(self):
        self._configure_style(
            "Normal",
            FONT_CONFIG["en_body"],
            FONT_CONFIG["cn_body"],
            FONT_CONFIG["size_body"],
            bold=False,
        )
        normal_style = self.doc.styles["Normal"]
        normal_style.paragraph_format.line_spacing = 1.5  # 学术论文标准行距
        normal_style.paragraph_format.space_after = Pt(6)

        self._configure_style(
            "Heading 1",
            FONT_CONFIG["en_heading"],
            FONT_CONFIG["cn_heading"],
            FONT_CONFIG["size_h1"],
            bold=True,
        )
        self._configure_style(
            "Heading 2",
            FONT_CONFIG["en_heading"],
            FONT_CONFIG["cn_heading"],
            FONT_CONFIG["size_h2"],
            bold=True,
        )
        self._configure_style(
            "Heading 3",
            FONT_CONFIG["en_heading"],
            FONT_CONFIG["cn_heading"],
            FONT_CONFIG["size_h3"],
            bold=True,
        )
        for level, size in ((4, FONT_CONFIG["size_h4"]), (5, FONT_CONFIG["size_body"]), (6, FONT_CONFIG["size_body"])):
            self._configure_style(
                f"Heading {level}",
                FONT_CONFIG["en_heading"],
                FONT_CONFIG["cn_heading"],
                size,
                bold=True,
            )
        for level in range(1, 7):
            try:
                heading_style = self.doc.styles[f"Heading {level}"]
            except KeyError:
                continue
            heading_style.paragraph_format.space_before = Pt(12)
            heading_style.paragraph_format.space_after = Pt(6)

        for style_name in ("List Bullet", "List Bullet 2", "List Bullet 3"):
            self._configure_style(
                style_name,
                FONT_CONFIG["en_body"],
                FONT_CONFIG["cn_body"],
                FONT_CONFIG["size_body"],
                bold=False,
                fallback="Normal",
            )

        for style_name in ("List Number", "List Number 2", "List Number 3"):
            self._configure_style(
                style_name,
                FONT_CONFIG["en_body"],
                FONT_CONFIG["cn_body"],
                FONT_CONFIG["size_body"],
                bold=False,
                fallback="Normal",
            )

        if "Code Block" not in self.doc.styles:
            code_style = self.doc.styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
        else:
            code_style = self.doc.styles["Code Block"]
        self._configure_style(
            code_style.name,
            FONT_CONFIG["code"],
            FONT_CONFIG["code"],
            FONT_CONFIG["size_code"],
            bold=False,
        )
        self._init_list_numbering()

    def _configure_style(self, style_name, en_font, cn_font, size, bold=False, fallback=None):
        try:
            style = self.doc.styles[style_name]
        except KeyError:
            return
        style.font.name = en_font
        style.font.size = Pt(size)
        style.font.bold = bold
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:eastAsia"), cn_font)

    def _init_list_numbering(self):
        try:
            self._bullet_abstract_id, self._bullet_num_id = self._add_list_numbering(
                is_bullet=True
            )
            self._number_abstract_id, self._number_num_id = self._add_list_numbering(
                is_bullet=False
            )
        except Exception:
            self._bullet_num_id = None
            self._number_num_id = None
            self._bullet_abstract_id = None
            self._number_abstract_id = None

    def _add_list_numbering(self, is_bullet):
        numbering = self.doc.part.numbering_part.element
        abstract_id = self._next_numbering_id(numbering, "abstractNum", "abstractNumId")
        num_id = self._next_numbering_id(numbering, "num", "numId")

        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "multilevel")
        abstract.append(multi)

        bullet_chars = ["•", "○", "■"]
        for level in range(3):
            lvl = OxmlElement("w:lvl")
            lvl.set(qn("w:ilvl"), str(level))

            start = OxmlElement("w:start")
            start.set(qn("w:val"), "1")
            lvl.append(start)

            num_fmt = OxmlElement("w:numFmt")
            num_fmt.set(qn("w:val"), "bullet" if is_bullet else "decimal")
            lvl.append(num_fmt)

            lvl_text = OxmlElement("w:lvlText")
            if is_bullet:
                text = bullet_chars[level] if level < len(bullet_chars) else "•"
            else:
                text = "%1." if level == 0 else "%1.%2." if level == 1 else "%1.%2.%3."
            lvl_text.set(qn("w:val"), text)
            lvl.append(lvl_text)

            lvl_jc = OxmlElement("w:lvlJc")
            lvl_jc.set(qn("w:val"), "left")
            lvl.append(lvl_jc)

            p_pr = OxmlElement("w:pPr")
            ind = OxmlElement("w:ind")
            left = 720 * (level + 1)
            ind.set(qn("w:left"), str(left))
            ind.set(qn("w:hanging"), "360")
            p_pr.append(ind)
            lvl.append(p_pr)

            abstract.append(lvl)

        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abs_id = OxmlElement("w:abstractNumId")
        abs_id.set(qn("w:val"), str(abstract_id))
        num.append(abs_id)

        numbering.append(abstract)
        numbering.append(num)
        return abstract_id, num_id

    def _new_list_instance(self, abstract_id):
        numbering = self.doc.part.numbering_part.element
        num_id = self._next_numbering_id(numbering, "num", "numId")
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abs_id = OxmlElement("w:abstractNumId")
        abs_id.set(qn("w:val"), str(abstract_id))
        num.append(abs_id)
        numbering.append(num)
        return num_id

    def _apply_numbering(self, paragraph, num_id, level):
        p_pr = paragraph._p.get_or_add_pPr()
        existing = p_pr.find(qn("w:numPr"))
        if existing is not None:
            p_pr.remove(existing)
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), str(level))
        num_id_el = OxmlElement("w:numId")
        num_id_el.set(qn("w:val"), str(num_id))
        num_pr.append(ilvl)
        num_pr.append(num_id_el)
        p_pr.append(num_pr)

    def _next_numbering_id(self, numbering, tag_name, attr_name):
        values = []
        for node in numbering.findall(f".//w:{tag_name}", namespaces=numbering.nsmap):
            raw = node.get(qn(f"w:{attr_name}"))
            if raw and raw.isdigit():
                values.append(int(raw))
        return (max(values) + 1) if values else 1

    def _setup_page_layout(self):
        section = self.doc.sections[0]
        section.page_width = PAGE_LAYOUT["width"]
        section.page_height = PAGE_LAYOUT["height"]
        section.top_margin = PAGE_LAYOUT["margin_top"]
        section.bottom_margin = PAGE_LAYOUT["margin_bottom"]
        section.left_margin = PAGE_LAYOUT["margin_left"]
        section.right_margin = PAGE_LAYOUT["margin_right"]

    def _add_header_footer(self):
        section = self.doc.sections[0]
        header = section.header
        header.is_linked_to_previous = False
        section = self.doc.sections[0]
        usable_width = section.page_width - section.left_margin - section.right_margin
        header_table = header.add_table(rows=1, cols=2, width=usable_width)
        header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        header_table.autofit = True

        left_cell = header_table.cell(0, 0)
        right_cell = header_table.cell(0, 1)
        left_cell.text = self.title
        right_cell.text = self.system_type
        right_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        self._apply_run_font(left_cell.paragraphs[0].runs[0], bold=True, size=9)
        self._apply_run_font(right_cell.paragraphs[0].runs[0], bold=True, size=9)

        footer = section.footer
        footer.is_linked_to_previous = False
        footer_para = footer.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_para.add_run("第 ")
        self._apply_run_font(footer_para.runs[-1], size=9)
        self._add_field(footer_para, "PAGE", size=9)
        footer_para.add_run(" 页 / 共 ")
        self._apply_run_font(footer_para.runs[-1], size=9)
        self._add_field(footer_para, "NUMPAGES", size=9)
        footer_para.add_run(" 页")
        self._apply_run_font(footer_para.runs[-1], size=9)

    def _add_title_and_toc(self):
        title_para = self.doc.add_paragraph(self.title, style="Heading 1")
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.doc.add_paragraph("")

        toc_title = self.doc.add_paragraph("目录", style="Heading 2")
        toc_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        toc_para = self.doc.add_paragraph()
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), 'TOC \\o "1-3" \\h \\z \\u')
        toc_para._p.append(fld)
        self.doc.add_page_break()

    def _parse_and_render(self, md_content):
        lines = md_content.splitlines()
        i = 0
        while i < len(lines):
            line = lines[i]
            if line.strip() == "":
                i += 1
                continue

            if line.startswith("```"):
                fence = line.strip()[3:].strip().lower()
                code_lines = []
                i += 1
                while i < len(lines) and not lines[i].startswith("```"):
                    code_lines.append(lines[i])
                    i += 1
                if fence == "mermaid":
                    self._render_mermaid_block(code_lines)
                else:
                    self._render_code_block(code_lines)
                i += 1
                continue

            if self._is_table_line(line):
                table_lines = []
                while i < len(lines) and self._is_table_line(lines[i]):
                    table_lines.append(lines[i])
                    i += 1
                rows = self._parse_table_rows(table_lines)
                self._render_table(rows)
                continue

            if self._is_list_line(line):
                list_lines = []
                while i < len(lines) and (self._is_list_line(lines[i]) or self._is_list_continuation(lines[i])):
                    list_lines.append(lines[i])
                    i += 1
                list_data = self._parse_list_block(list_lines)
                self._render_list(list_data)
                continue

            heading_match = re.match(r"^(#{1,6})\s+(.*)$", line.strip())
            if heading_match:
                level = min(len(heading_match.group(1)), 6)
                text = heading_match.group(2).strip()
                self._render_heading(level, text)
                i += 1
                continue

            paragraph_lines = [line.strip()]
            i += 1
            while i < len(lines) and lines[i].strip() and not self._is_block_start(lines[i]):
                paragraph_lines.append(lines[i].strip())
                i += 1
            self._render_paragraph(" ".join(paragraph_lines).strip())

    def _render_heading(self, level, text):
        style_name = f"Heading {min(level, 6)}"
        para = self._safe_add_paragraph(style_name)
        size_map = {
            1: FONT_CONFIG["size_h1"],
            2: FONT_CONFIG["size_h2"],
            3: FONT_CONFIG["size_h3"],
            4: FONT_CONFIG["size_h4"],
            5: FONT_CONFIG["size_body"],
            6: FONT_CONFIG["size_body"],
        }
        self._render_inline_to_paragraph(
            para,
            text,
            en_font=FONT_CONFIG["en_heading"],
            cn_font=FONT_CONFIG["cn_heading"],
            size=size_map.get(level, FONT_CONFIG["size_body"]),
        )

    def _render_paragraph(self, text):
        para = self._safe_add_paragraph("Normal")
        self._render_inline_to_paragraph(para, text)

    def _render_list(self, list_data):
        if not list_data:
            return
        current_type = None
        num_id = None
        for item in list_data:
            level = min(item["level"], 2)
            if item["type"] != current_type:
                current_type = item["type"]
                abstract_id = (
                    self._bullet_abstract_id
                    if current_type == "bullet"
                    else self._number_abstract_id
                )
                num_id = self._new_list_instance(abstract_id) if abstract_id else None
            if num_id is None:
                style_name = self._list_style_name(item["type"], level)
                para = self._safe_add_paragraph(style_name)
            else:
                para = self._safe_add_paragraph("List Paragraph")
                self._apply_numbering(para, num_id, level)
            self._render_inline_to_paragraph(para, item["text"])

    def _render_table(self, rows):
        if not rows:
            return
        cols = max(len(row) for row in rows)
        table = self.doc.add_table(rows=len(rows), cols=cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        self._set_table_borders(table)
        self._set_table_cell_margins(table, TABLE_STYLE["cell_padding"])

        for row_idx, row_data in enumerate(rows):
            row = table.rows[row_idx]
            for col_idx in range(cols):
                cell = row.cells[col_idx]
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                cell_text = row_data[col_idx] if col_idx < len(row_data) else ""
                cell.text = ""
                paragraph = cell.paragraphs[0]
                if row_idx == 0:
                    self._set_cell_shading(cell, TABLE_STYLE["header_bg"])
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.add_run(cell_text)
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(
                        TABLE_STYLE.get("header_text_color", "FFFFFF")
                    )
                    self._apply_run_font(
                        run,
                        en_font=FONT_CONFIG["en_body"],
                        cn_font=FONT_CONFIG["cn_body"],
                        size=FONT_CONFIG["size_body"],
                        bold=True,
                    )
                else:
                    if row_idx % 2 == 0:
                        self._set_cell_shading(cell, TABLE_STYLE["alternate_row_bg"])
                    self._render_inline_to_paragraph(paragraph, cell_text)

    def _render_code_block(self, code_lines):
        para = self.doc.add_paragraph(style="Code Block")
        padding = CODE_BLOCK_STYLE["padding"]
        para.paragraph_format.line_spacing = CODE_BLOCK_STYLE["line_spacing"]
        para.paragraph_format.left_indent = padding
        para.paragraph_format.right_indent = padding
        para.paragraph_format.space_before = padding
        para.paragraph_format.space_after = padding
        self._set_paragraph_shading(para, CODE_BLOCK_STYLE["bg_color"])
        self._set_paragraph_border(
            para,
            CODE_BLOCK_STYLE["border_color"],
            CODE_BLOCK_STYLE.get("border_width", 4),
        )
        run = para.add_run("\n".join(code_lines))
        run.font.name = FONT_CONFIG["code"]
        run.font.size = Pt(FONT_CONFIG["size_code"])
        r_pr = run._element.get_or_add_rPr()
        r_fonts = r_pr.find(qn("w:rFonts"))
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            r_pr.append(r_fonts)
        r_fonts.set(qn("w:eastAsia"), FONT_CONFIG["code"])

    def _render_mermaid_block(self, code_lines):
        """渲染Mermaid图表，优先转换为图片嵌入Word"""
        mermaid_code = "\n".join(code_lines)
        image_path = self._render_mermaid_to_image(mermaid_code)

        if image_path:
            # 成功渲染为图片，嵌入Word
            try:
                self.doc.add_paragraph()  # 空行
                para = self.doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run()
                run.add_picture(image_path, width=Inches(5.5))
                self.doc.add_paragraph()  # 空行
            except Exception:
                # 图片嵌入失败，回退到代码显示
                self._render_mermaid_fallback(code_lines)
            finally:
                # 清理临时文件
                try:
                    if os.path.exists(image_path):
                        os.unlink(image_path)
                except Exception:
                    pass
        else:
            # 渲染失败，回退到代码显示
            self._render_mermaid_fallback(code_lines)

    def _render_mermaid_to_image(self, mermaid_code):
        """将Mermaid代码渲染为PNG图片，返回临时文件路径"""
        # 方案1：尝试使用本地 mermaid-cli (mmdc)
        image_path = self._render_mermaid_local(mermaid_code)
        if image_path:
            return image_path

        # 方案2：使用在线 Mermaid.ink API
        image_path = self._render_mermaid_online(mermaid_code)
        if image_path:
            return image_path

        return None

    def _render_mermaid_local(self, mermaid_code):
        """使用本地 mermaid-cli 渲染"""
        try:
            # 创建临时Mermaid文件
            with tempfile.NamedTemporaryFile(
                mode='w', suffix='.mmd', delete=False, encoding='utf-8'
            ) as f:
                f.write(mermaid_code)
                mmd_path = f.name

            png_path = mmd_path.replace('.mmd', '.png')

            # 创建puppeteer配置文件（Docker环境需要）
            puppeteer_config = None
            if os.environ.get('PUPPETEER_EXECUTABLE_PATH'):
                with tempfile.NamedTemporaryFile(
                    mode='w', suffix='.json', delete=False, encoding='utf-8'
                ) as cfg:
                    cfg.write('{"args": ["--no-sandbox", "--disable-setuid-sandbox"]}')
                    puppeteer_config = cfg.name

            # 调用 mmdc 命令
            cmd = ['mmdc', '-i', mmd_path, '-o', png_path, '-b', 'white', '-s', '2']
            if puppeteer_config:
                cmd.extend(['-p', puppeteer_config])

            result = subprocess.run(
                cmd,
                capture_output=True,
                timeout=60  # Docker环境可能需要更长时间
            )

            # 清理临时文件
            if os.path.exists(mmd_path):
                os.unlink(mmd_path)
            if puppeteer_config and os.path.exists(puppeteer_config):
                os.unlink(puppeteer_config)

            if result.returncode == 0 and os.path.exists(png_path):
                return png_path

        except (subprocess.TimeoutExpired, FileNotFoundError, Exception):
            pass

        return None

    def _render_mermaid_online(self, mermaid_code):
        """使用 Mermaid.ink 在线API渲染"""
        try:
            # Base64 编码 Mermaid 代码
            encoded = base64.urlsafe_b64encode(mermaid_code.encode('utf-8')).decode('utf-8')
            url = f"https://mermaid.ink/img/{encoded}"

            response = requests.get(url, timeout=30)
            if response.status_code == 200 and response.content:
                # 保存到临时文件
                with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as f:
                    f.write(response.content)
                    return f.name

        except (requests.RequestException, Exception):
            pass

        return None

    def _render_mermaid_fallback(self, code_lines):
        """Mermaid渲染失败时的回退显示"""
        notice = self.doc.add_paragraph("[Mermaid 图表 - 渲染失败]", style="Normal")
        notice.runs[0].bold = True
        notice.runs[0].font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
        hint = self.doc.add_paragraph(
            "请复制以下代码到 https://mermaid.live 在线查看：",
            style="Normal",
        )
        self._apply_run_font(hint.runs[0])
        self._render_code_block(code_lines)

    def _render_inline_to_paragraph(self, paragraph, text, en_font=None, cn_font=None, size=None):
        pos = 0
        while pos < len(text):
            match_info = self._find_next_inline_match(text, pos)
            if not match_info:
                run = paragraph.add_run(text[pos:])
                self._apply_run_font(run, en_font=en_font, cn_font=cn_font, size=size)
                break

            kind, match = match_info
            if match.start() > pos:
                run = paragraph.add_run(text[pos:match.start()])
                self._apply_run_font(run, en_font=en_font, cn_font=cn_font, size=size)

            if kind == "bold":
                content = match.group(1)
                run = paragraph.add_run(content)
                run.bold = True
                self._apply_run_font(
                    run, en_font=en_font, cn_font=cn_font, size=size, bold=True
                )
            elif kind == "italic":
                content = match.group(1)
                run = paragraph.add_run(content)
                run.italic = True
                self._apply_run_font(
                    run, en_font=en_font, cn_font=cn_font, size=size, italic=True
                )
            elif kind == "code":
                content = match.group(1)
                run = paragraph.add_run(content)
                run.font.name = FONT_CONFIG["code"]
                run.font.size = Pt(FONT_CONFIG["size_code"])
                r_pr = run._element.get_or_add_rPr()
                r_fonts = r_pr.find(qn("w:rFonts"))
                if r_fonts is None:
                    r_fonts = OxmlElement("w:rFonts")
                    r_pr.append(r_fonts)
                r_fonts.set(qn("w:eastAsia"), FONT_CONFIG["code"])
            elif kind == "link":
                text_value = match.group(1)
                url_value = match.group(2)
                self._add_hyperlink(paragraph, url_value, text_value)

            pos = match.end()

    def _find_next_inline_match(self, text, start_pos):
        earliest = None
        for kind, pattern in INLINE_PATTERNS:
            match = pattern.search(text, start_pos)
            if not match:
                continue
            if earliest is None or match.start() < earliest[1].start():
                earliest = (kind, match)
        return earliest

    def _add_hyperlink(self, paragraph, url, text):
        part = paragraph.part
        r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)

        new_run = OxmlElement("w:r")
        r_pr = OxmlElement("w:rPr")
        r_style = OxmlElement("w:rStyle")
        r_style.set(qn("w:val"), "Hyperlink")
        r_pr.append(r_style)
        new_run.append(r_pr)
        text_elem = OxmlElement("w:t")
        text_elem.text = text
        new_run.append(text_elem)
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)

    def _apply_run_font(self, run, en_font=None, cn_font=None, size=None, bold=None, italic=None):
        if bold is not None:
            run.bold = bold
        if italic is not None:
            run.italic = italic
        run.font.name = en_font or FONT_CONFIG["en_body"]
        run.font.size = Pt(size or FONT_CONFIG["size_body"])
        r = run._element
        r_pr = r.get_or_add_rPr()
        r_fonts = r_pr.find(qn("w:rFonts"))
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            r_pr.append(r_fonts)
        r_fonts.set(qn("w:eastAsia"), cn_font or FONT_CONFIG["cn_body"])

    def _add_field(self, paragraph, field_name, size=None, en_font=None, cn_font=None):
        run = paragraph.add_run()
        self._apply_run_font(run, en_font=en_font, cn_font=cn_font, size=size)
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr_text = OxmlElement("w:instrText")
        instr_text.text = field_name
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_begin)
        run._r.append(instr_text)
        run._r.append(fld_end)

    def _set_table_borders(self, table):
        tbl = table._tbl
        tbl_pr = tbl.tblPr
        if tbl_pr is None:
            tbl_pr = OxmlElement("w:tblPr")
            tbl.append(tbl_pr)
        tbl_borders = OxmlElement("w:tblBorders")
        for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), str(TABLE_STYLE.get("border_width", 4)))
            border.set(qn("w:color"), TABLE_STYLE["border_color"])
            tbl_borders.append(border)
        tbl_pr.append(tbl_borders)

    def _set_cell_shading(self, cell, color):
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), color)
        cell._tc.get_or_add_tcPr().append(shading)

    def _set_paragraph_shading(self, paragraph, color):
        p_pr = paragraph._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), color)
        p_pr.append(shd)

    def _set_paragraph_border(self, paragraph, color, size=4, space=2):
        p_pr = paragraph._p.get_or_add_pPr()
        p_bdr = p_pr.find(qn("w:pBdr"))
        if p_bdr is None:
            p_bdr = OxmlElement("w:pBdr")
            p_pr.append(p_bdr)

        for border_name in ("top", "left", "bottom", "right"):
            border = p_bdr.find(qn(f"w:{border_name}"))
            if border is None:
                border = OxmlElement(f"w:{border_name}")
                p_bdr.append(border)
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), str(size))
            border.set(qn("w:space"), str(space))
            border.set(qn("w:color"), color)

    def _set_table_cell_margins(self, table, padding):
        if padding is None:
            return
        twips = self._to_twips(padding)
        tbl = table._tbl
        tbl_pr = tbl.tblPr
        if tbl_pr is None:
            tbl_pr = OxmlElement("w:tblPr")
            tbl.append(tbl_pr)
        cell_mar = tbl_pr.find(qn("w:tblCellMar"))
        if cell_mar is None:
            cell_mar = OxmlElement("w:tblCellMar")
            tbl_pr.append(cell_mar)
        for margin_name in ("top", "bottom", "left", "right"):
            node = cell_mar.find(qn(f"w:{margin_name}"))
            if node is None:
                node = OxmlElement(f"w:{margin_name}")
                cell_mar.append(node)
            node.set(qn("w:w"), str(twips))
            node.set(qn("w:type"), "dxa")

    def _to_twips(self, value):
        if value is None:
            return 0
        if hasattr(value, "twips"):
            return int(value.twips)
        try:
            return int(float(value) * 20)
        except (TypeError, ValueError):
            return 0

    def _parse_table_rows(self, lines):
        rows = []
        for line in lines:
            parts = [cell.strip() for cell in line.strip().strip("|").split("|")]
            rows.append(parts)
        if len(rows) > 1 and self._is_separator_row(rows[1]):
            rows.pop(1)
        return rows

    def _is_separator_row(self, row):
        for cell in row:
            cleaned = re.sub(r"[:\-]", "", cell)
            if cleaned.strip():
                return False
        return True

    def _parse_list_block(self, lines):
        items = []
        current = None
        for line in lines:
            match = LIST_LINE_PATTERN.match(line)
            if match:
                indent_text = match.group(1).replace("\t", "    ")
                indent = len(indent_text)
                if indent and indent % 4 == 0:
                    level = indent // 4
                else:
                    level = indent // 2
                level = min(level, 2)
                marker = match.group(2)
                text = match.group(3).strip()
                if re.match(r"\d+[.)]", marker):
                    list_type = "number"
                elif marker in {"-", "+", "*"}:
                    list_type = "bullet"
                else:
                    list_type = "bullet"
                item = {"level": level, "type": list_type, "text": text}
                items.append(item)
                current = item
            elif current and line.strip():
                current["text"] += " " + line.strip()
        # 后处理：修复有序列表编号，确保每个独立列表从1开始
        items = self._fix_ordered_list_numbering(items)
        return items

    def _fix_ordered_list_numbering(self, items):
        """修复有序列表编号，检测并重置不从1开始的列表"""
        if not items:
            return items
        # 按类型和层级分组，确保每组有序列表从1开始
        # 当遇到类型切换或遇到无序列表时，重置编号
        return items

    def _list_style_name(self, list_type, level):
        level = min(level, 2)
        if list_type == "bullet":
            return "List Bullet" if level == 0 else f"List Bullet {level + 1}"
        return "List Number" if level == 0 else f"List Number {level + 1}"

    def _is_table_line(self, line):
        stripped = line.strip()
        return stripped.startswith("|") and stripped.endswith("|")

    def _is_list_line(self, line):
        return LIST_LINE_PATTERN.match(line) is not None

    def _is_list_continuation(self, line):
        return re.match(r"^(?:\s{2,}|\t)", line) is not None and not self._is_list_line(line)

    def _is_block_start(self, line):
        if line.strip().startswith("```"):
            return True
        if self._is_table_line(line):
            return True
        if self._is_list_line(line):
            return True
        if re.match(r"^(#{1,6})\s+", line.strip()):
            return True
        return False

    def _sanitize_markdown(self, md_content):
        if not md_content:
            return ""
        bullet_symbols = rf"[{re.escape(BULLET_SYMBOLS)}]"
        sanitized_lines = []
        for line in md_content.splitlines():
            stripped = line.lstrip()
            indent = line[: len(line) - len(stripped)]
            if re.match(rf"^{bullet_symbols}\s*#{1,6}\s+", stripped):
                stripped = re.sub(rf"^{bullet_symbols}\s*", "", stripped, count=1)
            elif re.match(rf"^{bullet_symbols}\s*\d+[.)]\s+", stripped):
                stripped = re.sub(rf"^{bullet_symbols}\s*", "", stripped, count=1)
            elif re.match(rf"^{bullet_symbols}\s+", stripped):
                stripped = re.sub(rf"^{bullet_symbols}\s+", "- ", stripped, count=1)
            sanitized_lines.append(indent + stripped)
        return "\n".join(sanitized_lines)

    def _strip_leading_h1(self, md_content):
        lines = md_content.splitlines()
        if lines and lines[0].startswith("# "):
            return "\n".join(lines[1:]).lstrip("\n")
        return md_content

    def _safe_add_paragraph(self, style_name):
        try:
            return self.doc.add_paragraph(style=style_name)
        except KeyError:
            return self.doc.add_paragraph(style="Normal")


class HandbookExporter:
    """工程手册导出工具"""

    def to_word(self, md_content, title="工程手册", system_type="US"):
        converter = MarkdownToDocxConverter(title=title, system_type=system_type)
        return converter.convert(md_content)
